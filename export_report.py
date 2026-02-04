"""
Export Report - Convert generated reports to Word documents.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def is_section_header(line: str) -> bool:
    """
    Check if a line is a section header.

    Headers are identified as:
    - Lines ending with ':'
    - Lines that are all uppercase (with optional punctuation)
    - Lines starting with '**' (markdown bold)
    """
    line = line.strip()
    if not line:
        return False

    # Remove markdown bold markers for checking
    clean_line = line.replace('**', '').strip()

    # Check if line ends with colon
    if clean_line.endswith(':'):
        return True

    # Check if line is all uppercase (allowing spaces and punctuation)
    alpha_chars = ''.join(c for c in clean_line if c.isalpha())
    if alpha_chars and alpha_chars.isupper() and len(alpha_chars) >= 3:
        return True

    return False


def clean_markdown(text: str) -> str:
    """Remove markdown formatting from text."""
    # Remove bold markers
    text = text.replace('**', '')
    return text.strip()


def export_to_docx(report_text: str, filename: str) -> str:
    """
    Export report text to a Word document.

    Args:
        report_text: The generated operative report text
        filename: Base filename (without extension)

    Returns:
        Full path to the created .docx file
    """
    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Add title
    title = doc.add_heading('OPERATIVE REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a blank line after title
    doc.add_paragraph()

    # Split report into lines and process
    lines = report_text.split('\n')
    current_paragraph = []

    for line in lines:
        line = line.strip()

        # Skip empty lines - they indicate paragraph breaks
        if not line:
            # Flush current paragraph if any
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                doc.add_paragraph(clean_markdown(para_text))
                current_paragraph = []
            continue

        # Check if this is a section header
        if is_section_header(line):
            # Flush current paragraph first
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                doc.add_paragraph(clean_markdown(para_text))
                current_paragraph = []

            # Add header as bold paragraph
            clean_header = clean_markdown(line)
            para = doc.add_paragraph()
            run = para.add_run(clean_header)
            run.bold = True

        # Check if this is a list item
        elif line.startswith('- ') or line.startswith('• '):
            # Flush current paragraph first
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                doc.add_paragraph(clean_markdown(para_text))
                current_paragraph = []

            # Add as list item with indent
            clean_item = clean_markdown(line[2:])  # Remove bullet
            para = doc.add_paragraph(clean_item, style='List Bullet')

        # Check if this is a numbered list item
        elif re.match(r'^\d+\.?\s', line):
            # Flush current paragraph first
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                doc.add_paragraph(clean_markdown(para_text))
                current_paragraph = []

            # Add as numbered item
            clean_item = clean_markdown(re.sub(r'^\d+\.?\s*', '', line))
            para = doc.add_paragraph(clean_item, style='List Number')

        else:
            # Regular text - accumulate into paragraph
            current_paragraph.append(line)

    # Flush any remaining paragraph
    if current_paragraph:
        para_text = ' '.join(current_paragraph)
        doc.add_paragraph(clean_markdown(para_text))

    # Ensure filename doesn't have extension
    if filename.endswith('.docx'):
        filename = filename[:-5]

    # Save to /tmp
    filepath = f"/tmp/{filename}.docx"
    doc.save(filepath)

    return filepath


if __name__ == "__main__":
    # Test with sample report
    sample_report = """**OPERATIVE REPORT**

**PREOPERATIVE DIAGNOSIS:** Symptomatic cholelithiasis with biliary colic.

**POSTOPERATIVE DIAGNOSIS:** Symptomatic cholelithiasis with chronic cholecystitis.

**PROCEDURE PERFORMED:** Laparoscopic Cholecystectomy

**SURGEON:** Dr. Smith
**ASSISTANT:** Dr. Johnson

**ANESTHESIA:** General endotracheal anesthesia was induced and maintained throughout the procedure.

**INDICATIONS:** The patient is a 45-year-old female with recurrent right upper quadrant pain, confirmed by ultrasound to have gallstones. Conservative management has failed.

**FINDINGS:** Intraoperative findings revealed a distended gallbladder with omental adhesions. Multiple faceted stones were palpable within the gallbladder.

**PROCEDURE IN DETAIL:**
- The patient was brought to the operating room and anesthesia was induced.
- The abdomen was prepped and draped in the usual sterile fashion.
- A standard 4-port technique was utilized.
- Critical view of safety achieved.
- Cystic duct and artery doubly clipped and divided.
- Gallbladder dissected from liver bed using electrocautery.
- Gallbladder removed via umbilical port in specimen bag.

**SPECIMENS:** The gallbladder with stones was sent to pathology.

**DRAINS:** No drains were placed.

**ESTIMATED BLOOD LOSS:** Less than 20 mL

**COMPLICATIONS:** None noted during or at the conclusion of the procedure.

**DISPOSITION:** The patient tolerated the surgery well and was transferred to the recovery room in satisfactory condition."""

    print("Export Report - Test")
    print("=" * 60)

    filepath = export_to_docx(sample_report, "test_operative_report")

    print(f"\nDocument created: {filepath}")
    print(f"File exists: {os.path.exists(filepath)}")
    print(f"File size: {os.path.getsize(filepath)} bytes")
